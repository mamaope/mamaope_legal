import os
import io
import re
import boto3
import argparse
import time
import logging
from dotenv import load_dotenv
from pymilvus import MilvusClient, DataType
import openai
import fitz
import docx
import pandas as pd
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type, before_sleep_log

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    handlers=[logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

# Suppress MuPDF warnings
logging.getLogger("fitz").setLevel(logging.WARNING)

# Load environment variables
load_dotenv()

# Configuration
S3_BUCKET_NAME = os.getenv('AWS_S3_BUCKET', 'mamaope-legal')
S3_INPUT_PREFIX = 'input/'
MILVUS_URI = os.getenv('MILVUS_URI')
MILVUS_TOKEN = os.getenv('MILVUS_TOKEN')
COLLECTION_NAME = os.getenv('MILVUS_COLLECTION_NAME', 'mamaope_legal')

# Initialize Azure OpenAI client
azure_client = openai.AzureOpenAI(
    azure_endpoint=os.getenv('AZURE_OPENAI_ENDPOINT'),
    api_key=os.getenv('AZURE_OPENAI_API_KEY'),
    api_version=os.getenv('API_VERSION', '2024-02-01')
)
AZURE_DEPLOYMENT = os.getenv('DEPLOYMENT', 'text-embedding-3-large')

def initialize_zilliz_collection(refresh: bool = False):
    """
    Connects to Milvus and initializes the collection with a schema for Azure embeddings.
    """
    setup_client = MilvusClient(uri=MILVUS_URI, token=MILVUS_TOKEN)
    
    try:
        if refresh and setup_client.has_collection(collection_name=COLLECTION_NAME):
            logger.info(f"Refresh mode: Dropping existing collection '{COLLECTION_NAME}'...")
            setup_client.drop_collection(collection_name=COLLECTION_NAME)
            logger.info("Collection dropped.")

        if not setup_client.has_collection(collection_name=COLLECTION_NAME):
            logger.info(f"Creating collection '{COLLECTION_NAME}'...")
            schema = MilvusClient.create_schema(auto_id=True, enable_dynamic_field=False)
            schema.add_field(field_name="id", datatype=DataType.INT64, is_primary=True)
            schema.add_field(field_name="content", datatype=DataType.VARCHAR, max_length=65535)
            schema.add_field(field_name="file_path", datatype=DataType.VARCHAR, max_length=1024, is_partition_key=True)
            schema.add_field(field_name="display_page_number", datatype=DataType.VARCHAR, max_length=100)        
            
            schema.add_field(field_name="vector", datatype=DataType.FLOAT_VECTOR, dim=3072, is_nullable=False)
            index_params = MilvusClient.prepare_index_params()
            index_params.add_index(field_name="vector", index_type="AUTOINDEX", metric_type="COSINE")
            
            setup_client.create_collection(
                collection_name=COLLECTION_NAME,
                schema=schema,
                index_params=index_params
            )
            logger.info(f"Collection '{COLLECTION_NAME}' created successfully.")
            logger.info("Waiting 5 seconds for collection to initialize...")
            time.sleep(5)
        else:
            logger.info(f"Collection '{COLLECTION_NAME}' already exists. Proceeding with update.")
    
    finally:
        setup_client.close()

def extract_content_with_pymupdf_semantic(pdf_content: bytes, file_key: str) -> list[dict]:
    """
    Extracts semantic chunks from a PDF using PyMuPDF.
    """
    logger.info(f"Parsing PDF: {file_key}...")
    chunks = []
    
    try:
        with fitz.open(stream=pdf_content, filetype="pdf") as doc:
            for page_num, page in enumerate(doc, start=1):
                text = page.get_text("text")
                paragraphs = re.split(r'\n\s*\n', text)

                for para in paragraphs:
                    if not para.strip():
                        continue
                    words = para.split()
                    word_count = len(words)
                    
                    if 10 < word_count < 400:
                        chunks.append({'text': para.strip(), 'page': page_num})
                    elif word_count >= 400:
                        chunk_size, overlap = 300, 50
                        for i in range(0, word_count, chunk_size - overlap):
                            chunk_text = " ".join(words[i : i + chunk_size])
                            chunks.append({'text': chunk_text, 'page': page_num})

                tables = page.find_tables()
                if tables:
                    for i, table in enumerate(tables):
                        try:
                            table_data = table.extract()
                            table_text = "\n".join([" | ".join([str(cell) if cell is not None else "" for cell in row]) for row in table_data])
                            content = f"The following is data from a table on page {page_num}:\n{table_text}"
                            chunks.append({'text': content, 'page': page_num})
                        except Exception as e:
                            logger.warning(f"Could not extract table {i} on page {page_num}. Error: {e}")

        logger.info(f"Extracted {len(chunks)} chunks from {file_key}.")
        return chunks
    except Exception as e:
        logger.error(f"Failed to parse PDF {file_key}: {e}")
        return []

def extract_content_with_docx(docx_content: bytes, file_key: str) -> list[dict]:
    """
    Extracts semantic chunks from a DOCX file using python-docx.
    """
    logger.info(f"Parsing DOCX: {file_key}...")
    chunks = []
    
    try:
        doc = docx.Document(io.BytesIO(docx_content))
        section_number = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            words = text.split()
            word_count = len(words)
            section_number += 1
            
            if 10 < word_count < 400:
                chunks.append({'text': text, 'page': section_number})
            elif word_count >= 400:
                chunk_size, overlap = 300, 50
                for i in range(0, word_count, chunk_size - overlap):
                    chunk_text = " ".join(words[i : i + chunk_size])
                    chunks.append({'text': chunk_text, 'page': section_number})

        for table in doc.tables:
            section_number += 1
            try:
                table_text = "\n".join([" | ".join([cell.text.strip() if cell.text else "" for cell in row.cells]) for row in table.rows])
                content = f"The following is data from a table in section {section_number}:\n{table_text}"
                chunks.append({'text': content, 'page': section_number})
            except Exception as e:
                logger.warning(f"Could not extract table in section {section_number}. Error: {e}")

        logger.info(f"Extracted {len(chunks)} chunks from {file_key}.")
        return chunks
    except Exception as e:
        logger.error(f"Failed to parse DOCX {file_key}: {e}")
        return []

def extract_content_with_txt(txt_content: bytes, file_key: str) -> list[dict]:
    """
    Extracts semantic chunks from a TXT file.
    """
    logger.info(f"Parsing TXT: {file_key}...")
    chunks = []
    
    try:
        text = txt_content.decode('utf-8', errors='ignore')
        paragraphs = re.split(r'\n\s*\n', text)
        section_number = 0
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            words = para.split()
            word_count = len(words)
            section_number += 1
            
            if 10 < word_count < 400:
                chunks.append({'text': para, 'page': section_number})
            elif word_count >= 400:
                chunk_size, overlap = 300, 50
                for i in range(0, word_count, chunk_size - overlap):
                    chunk_text = " ".join(words[i : i + chunk_size])
                    chunks.append({'text': chunk_text, 'page': section_number})

        logger.info(f"Extracted {len(chunks)} chunks from {file_key}.")
        return chunks
    except Exception as e:
        logger.error(f"Failed to parse TXT {file_key}: {e}")
        return []

def extract_content_with_csv(csv_content: bytes, file_key: str) -> list[dict]:
    """
    Extracts chunks from a CSV file using pandas, treating each row as a chunk.
    """
    logger.info(f"Parsing CSV: {file_key}...")
    chunks = []
    
    try:
        df = pd.read_csv(io.BytesIO(csv_content), encoding='utf-8', errors='ignore')
        for idx, row in df.iterrows():
            row_text = " | ".join([str(val) if pd.notnull(val) else "" for val in row])
            content = f"CSV row {idx + 1}:\n{row_text}"
            chunks.append({'text': content, 'page': idx + 1})

        logger.info(f"Extracted {len(chunks)} chunks from {file_key}.")
        return chunks
    except Exception as e:
        logger.error(f"Failed to parse CSV {file_key}: {e}")
        return []

def extract_content(file_key: str, content: bytes) -> list[dict]:
    """
    Dispatches content extraction based on file extension.
    """
    file_extension = os.path.splitext(file_key)[1].lower()
    if file_extension == '.pdf':
        return extract_content_with_pymupdf_semantic(content, file_key)
    elif file_extension == '.docx':
        return extract_content_with_docx(content, file_key)
    elif file_extension == '.txt':
        return extract_content_with_txt(content, file_key)
    elif file_extension == '.csv':
        return extract_content_with_csv(content, file_key)
    else:
        logger.warning(f"Unsupported file type: {file_extension} for {file_key}")
        return []

@retry(
    stop=stop_after_attempt(5),
    wait=wait_exponential(multiplier=1, min=4, max=60),
    retry=retry_if_exception_type(openai.RateLimitError),
    before_sleep=before_sleep_log(logger, logging.INFO)
)
def generate_embeddings_batch(texts: list[str]) -> list[list[float]]:
    """
    Generates embeddings for a batch of texts using Azure OpenAI.
    """
    logger.info(f"Generating Azure OpenAI embeddings for {len(texts)} texts using deployment '{AZURE_DEPLOYMENT}'...")
    try:
        response = azure_client.embeddings.create(
            input=texts,
            model=AZURE_DEPLOYMENT
        )
        embeddings = [item.embedding for item in response.data]
        logger.info(f"Successfully generated {len(embeddings)} embeddings.")
        return embeddings
    except openai.RateLimitError as e:
        retry_after = e.response.headers.get('Retry-After', 60)
        logger.info(f"Rate limit hit. Retrying after {retry_after} seconds.")
        raise
    except Exception as e:
        logger.error(f"Embedding generation failed: {e}")
        raise

def generate_embeddings(texts: list[str], batch_size: int = 100) -> list[list[float]]:
    """
    Generates embeddings in batches to respect Azure rate limits.
    """
    all_embeddings = []
    for i in range(0, len(texts), batch_size):
        batch_texts = texts[i:i + batch_size]
        logger.info(f"Processing embedding batch {i//batch_size + 1}/{len(texts)//batch_size + 1} ({len(batch_texts)} texts)...")
        embeddings = generate_embeddings_batch(batch_texts)
        all_embeddings.extend(embeddings)
        time.sleep(1)
    return all_embeddings

def main(refresh: bool, specific_file: str = None):
    """Main function to run the ingestion process."""
    if refresh:
        logger.info("--- Starting Data Ingestion in FULL REFRESH mode ---")
    else:
        logger.info("--- Starting Data Ingestion in INCREMENTAL UPDATE mode ---")

    # Initialize collection
    initialize_zilliz_collection(refresh=refresh)

    # Create fresh Milvus client for insertion
    logger.info("Creating a fresh client for data insertion...")
    insert_client = MilvusClient(uri=MILVUS_URI, token=MILVUS_TOKEN)

    # Initialize S3 client
    s3 = boto3.client(
        's3',
        aws_access_key_id=os.getenv('AWS_ACCESS_KEY_ID'),
        aws_secret_access_key=os.getenv('AWS_SECRET_ACCESS_KEY'),
        region_name=os.getenv('AWS_REGION')
    )

    try:
        # List files in S3 bucket
        response = s3.list_objects_v1(Bucket=S3_BUCKET_NAME, Prefix=S3_INPUT_PREFIX)
        supported_extensions = {'.pdf', '.docx', '.txt', '.csv'}
        all_s3_files = [
            obj['Key'] for obj in response.get('Contents', [])
            if os.path.splitext(obj['Key'])[1].lower() in supported_extensions
        ]
        if specific_file:
            all_s3_files = [specific_file] if specific_file in all_s3_files else []
        
        logger.info(f"Found {len(all_s3_files)} supported files to process: {all_s3_files}")
        
        total_chunks_ingested = 0
        for file_key in all_s3_files:
            logger.info(f"Processing file: {file_key}...")
            try:
                # Download file from S3
                file_obj = s3.get_object(Bucket=S3_BUCKET_NAME, Key=file_key)
                file_content = file_obj['Body'].read()
                
                # Extract chunks
                chunks = extract_content(file_key, file_content)
                if not chunks:
                    logger.warning(f"No content extracted from {file_key}. Skipping.")
                    continue
                
                # Generate embeddings in batches
                texts = [chunk['text'] for chunk in chunks]
                embeddings = generate_embeddings(texts, batch_size=100)
                
                # Prepare data for insertion
                data_to_insert = [
                    {
                        "content": chunk['text'], 
                        "file_path": file_key, 
                        "display_page_number": str(chunk['page']),
                        "vector": emb
                    }
                    for chunk, emb in zip(chunks, embeddings)
                ]
                
                # Insert in batches
                batch_size = 100
                for i in range(0, len(data_to_insert), batch_size):
                    batch = data_to_insert[i:i + batch_size]
                    try:
                        res = insert_client.insert(collection_name=COLLECTION_NAME, data=batch)
                        total_chunks_ingested += res['insert_count']
                        logger.info(f"Inserted batch {i//batch_size + 1} ({len(batch)} chunks) for {file_key}")
                    except Exception as e:
                        logger.error(f"Failed to insert batch {i//batch_size + 1} for {file_key}: {e}")
                
                logger.info(f"Successfully inserted {len(data_to_insert)} chunks for {file_key}.")

            except Exception as e:
                logger.error(f"Failed to process file {file_key}: {e}")
                continue
        
        logger.info("--- Ingestion Complete ---")
        logger.info(f"Total chunks ingested in this run: {total_chunks_ingested}")
        
    except Exception as e:
        logger.error(f"Error listing S3 files: {e}")
    
    finally:
        insert_client.close()

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Run the Zilliz ingestion script for medical knowledge base.")
    parser.add_argument(
        '--refresh',
        action='store_true',
        help='If set, drops the existing collection and re-ingests all documents.'
    )
    parser.add_argument(
        '--file',
        type=str,
        default=None,
        help='Specific file to process (e.g., input/document.pdf).'
    )
    args = parser.parse_args()
    main(refresh=args.refresh, specific_file=args.file)
    